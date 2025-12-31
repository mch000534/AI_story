"""
AI Story Backend - Export Service
"""
import os
import io
import json
import zipfile
from datetime import datetime
from typing import Optional

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side

from app.models import Project, Stage, StageType, STAGE_NAMES


class ExportService:
    """Service for exporting project content."""
    
    def export_script_pdf(self, project: Project, stages: list[Stage]) -> bytes:
        """Export script stage as PDF."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        # Register Chinese Font
        # Use absolute path relative to the file
        base_dir = os.path.dirname(os.path.abspath(__file__))
        font_path = os.path.join(os.path.dirname(base_dir), 'static', 'fonts', 'NotoSansTC-Regular.ttf')
        
        try:
            pdfmetrics.registerFont(TTFont('NotoSansTC', font_path))
            font_name = 'NotoSansTC'
        except Exception as e:
            # Fallback to Helvetica only if font file is missing, otherwise raise to debug
            print(f"Failed to load font from {font_path}: {e}")
            if os.path.exists(font_path):
                raise RuntimeError(f"Font file exists but failed to load: {e}")
            font_name = 'Helvetica' # Still fallback if file missing to avoid crash, but log it.
            # actually, if we fallback, we get tofu. Better to fail?
            # Let's fail if we are sure we need Chinese.
            # But the user might be exporting English content.
            # For now, let's keep fallback but ensure path is correct.
            # The issue user sees implies fallback happened or crash happened.
            # If 500, crash happened.
            # If fallback happened, no 500, but tofu.
            # User reported 500. So crash happened.
            # If crash happened, it might be inside registerFont or later.
            # The previous 'print' swallowed the error.
            # So the 500 comes from later 'UnicodeEncodeError' or similar.
            pass

        story = []
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Title'],
            fontName=font_name,
            fontSize=24,
            spaceAfter=30
        )
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading1'],
            fontName=font_name,
            fontSize=16,
            spaceBefore=20,
            spaceAfter=10
        )
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=11,
            leading=16
        )
        
        # Add normal style with Chinese font for simple paragraphs if needed
        styles['Normal'].fontName = font_name
        
        # Title
        story.append(Paragraph(project.name, title_style))
        story.append(Paragraph(f"生成時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
        story.append(Spacer(1, 30))
        
        # Add each stage content
        for stage in stages:
            if stage.content:
                stage_name = STAGE_NAMES.get(stage.stage_type, stage.stage_type.value)
                story.append(Paragraph(stage_name, heading_style))
                
                # Split content into paragraphs
                paragraphs = stage.content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        # Escape special characters
                        safe_para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(safe_para, body_style))
                        story.append(Spacer(1, 6))
                
                story.append(PageBreak())
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_script_docx(self, project: Project, stages: list[Stage]) -> bytes:
        """Export script as Word document."""
        doc = Document()
        
        # Title
        title = doc.add_heading(project.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"生成時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()
        
        # Add each stage content
        for stage in stages:
            if stage.content:
                stage_name = STAGE_NAMES.get(stage.stage_type, stage.stage_type.value)
                doc.add_heading(stage_name, 1)
                
                paragraphs = stage.content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para)
                
                doc.add_page_break()
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_fountain(self, project: Project, script_stage: Optional[Stage]) -> str:
        """Export script as Fountain format."""
        if not script_stage or not script_stage.content:
            return ""
        
        # Basic Fountain conversion
        lines = []
        lines.append(f"Title: {project.name}")
        lines.append(f"Credit: 由 AI 故事創作工具生成")
        lines.append(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        lines.append("")
        lines.append("===")
        lines.append("")
        
        # Add script content (basic formatting)
        content = script_stage.content
        lines.append(content)
        
        return '\n'.join(lines)
    
    def export_storyboard_excel(self, project: Project, storyboard_stage: Optional[Stage]) -> bytes:
        """Export storyboard as Excel spreadsheet."""
        wb = Workbook()
        ws = wb.active
        ws.title = "分鏡表"
        
        # Headers
        headers = ['鏡號', '景別', '運鏡', '畫面描述', '對白/音效', '時長', '備註']
        header_font = Font(bold=True, size=12)
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = thin_border
        
        # Column widths
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 10
        ws.column_dimensions['C'].width = 10
        ws.column_dimensions['D'].width = 40
        ws.column_dimensions['E'].width = 30
        ws.column_dimensions['F'].width = 10
        ws.column_dimensions['G'].width = 20
        
        # Parse storyboard content and add rows
        if storyboard_stage and storyboard_stage.content:
            row = 2
            # Simple parsing - each line as a shot
            lines = storyboard_stage.content.split('\n')
            shot_num = 1
            for line in lines:
                if line.strip() and not line.startswith('#'):
                    ws.cell(row=row, column=1, value=shot_num).border = thin_border
                    ws.cell(row=row, column=4, value=line.strip()).border = thin_border
                    for col in [2, 3, 5, 6, 7]:
                        ws.cell(row=row, column=col).border = thin_border
                    shot_num += 1
                    row += 1
        
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_prompts_txt(self, project: Project, image_stage: Optional[Stage], motion_stage: Optional[Stage]) -> str:
        """Export AI prompts as text file."""
        lines = []
        lines.append(f"# {project.name} - AI 提示詞")
        lines.append(f"# 生成時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        lines.append("")
        
        if image_stage and image_stage.content:
            lines.append("## AI 圖像提示詞")
            lines.append("")
            lines.append(image_stage.content)
            lines.append("")
        
        if motion_stage and motion_stage.content:
            lines.append("## 動態分鏡提示詞")
            lines.append("")
            lines.append(motion_stage.content)
        
        return '\n'.join(lines)
    
    def export_complete_zip(self, project: Project, stages: list[Stage]) -> bytes:
        """Export complete project as ZIP archive."""
        buffer = io.BytesIO()
        
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            # Project info
            project_info = {
                "name": project.name,
                "description": project.description,
                "created_at": project.created_at.isoformat(),
                "exported_at": datetime.now().isoformat()
            }
            zf.writestr("project_info.json", json.dumps(project_info, ensure_ascii=False, indent=2))
            
            # Each stage as separate file
            for stage in stages:
                if stage.content:
                    stage_name = STAGE_NAMES.get(stage.stage_type, stage.stage_type.value)
                    filename = f"{stage.stage_type.value}.md"
                    content = f"# {stage_name}\n\n{stage.content}"
                    zf.writestr(filename, content.encode('utf-8'))
            
            # PDF (if script exists)
            script_stage = next((s for s in stages if s.stage_type == StageType.SCRIPT and s.content), None)
            if script_stage:
                try:
                    pdf_content = self.export_script_pdf(project, stages)
                    zf.writestr("script.pdf", pdf_content)
                except Exception:
                    pass
            
            # Excel (if storyboard exists)
            storyboard_stage = next((s for s in stages if s.stage_type == StageType.STORYBOARD and s.content), None)
            if storyboard_stage:
                try:
                    excel_content = self.export_storyboard_excel(project, storyboard_stage)
                    zf.writestr("storyboard.xlsx", excel_content)
                except Exception:
                    pass
            
            # Prompts
            image_stage = next((s for s in stages if s.stage_type == StageType.IMAGE_PROMPT and s.content), None)
            motion_stage = next((s for s in stages if s.stage_type == StageType.MOTION_PROMPT and s.content), None)
            if image_stage or motion_stage:
                prompts_content = self.export_prompts_txt(project, image_stage, motion_stage)
                zf.writestr("prompts.txt", prompts_content.encode('utf-8'))
        
        buffer.seek(0)
        return buffer.getvalue()
